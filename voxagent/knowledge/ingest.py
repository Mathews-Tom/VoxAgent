from __future__ import annotations

import hashlib
import re
from collections import deque
from dataclasses import dataclass
from urllib.parse import urljoin, urlparse

import httpx
from lxml import html as lxml_html
from readability import Document


@dataclass
class PageContent:
    url: str
    title: str
    html: str
    text: str
    content_hash: str
    source_type: str = "web"
    source_version_id: str | None = None


def _compute_hash(text: str) -> str:
    return hashlib.sha256(text.encode()).hexdigest()


def _extract_links(base_url: str, html_content: str) -> list[str]:
    base_parsed = urlparse(base_url)
    tree = lxml_html.fromstring(html_content, base_url=base_url)
    links: list[str] = []
    for anchor in tree.xpath("//a[@href]"):
        href = anchor.get("href", "").strip()
        if not href or href.startswith(("#", "mailto:", "tel:", "javascript:")):
            continue
        absolute = urljoin(base_url, href)
        parsed = urlparse(absolute)
        if parsed.hostname == base_parsed.hostname and parsed.scheme in ("http", "https"):
            # Strip fragments
            clean = parsed._replace(fragment="").geturl()
            links.append(clean)
    return links


def _parse_page(url: str, raw_html: str) -> PageContent:
    doc = Document(raw_html)
    title = doc.title() or ""
    clean_html = doc.summary()
    tree = lxml_html.fromstring(clean_html)
    text = tree.text_content()
    # Collapse excessive whitespace
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    content_hash = _compute_hash(text)
    return PageContent(url=url, title=title, html=raw_html, text=text, content_hash=content_hash)


async def crawl_website(
    base_url: str,
    max_depth: int = 3,
    max_pages: int = 100,
) -> list[PageContent]:
    visited: set[str] = set()
    results: list[PageContent] = []
    # Queue entries: (url, depth)
    queue: deque[tuple[str, int]] = deque([(base_url, 0)])

    async with httpx.AsyncClient(
        follow_redirects=True,
        timeout=httpx.Timeout(30.0),
        headers={"User-Agent": "VoxAgent-Crawler/1.0"},
    ) as client:
        while queue and len(results) < max_pages:
            url, depth = queue.popleft()
            if url in visited:
                continue
            visited.add(url)

            try:
                response = await client.get(url)
            except (httpx.RequestError, httpx.HTTPStatusError):
                continue

            if response.status_code != 200:
                continue

            content_type = response.headers.get("content-type", "")
            if "text/html" not in content_type:
                continue

            raw_html = response.text
            page = _parse_page(url, raw_html)
            results.append(page)

            if depth < max_depth:
                for link in _extract_links(url, raw_html):
                    if link not in visited:
                        queue.append((link, depth + 1))

    return results


def extract_text_from_file(file_path: str) -> str:
    if file_path.endswith(".txt"):
        with open(file_path, encoding="utf-8") as fh:
            return fh.read()

    if file_path.endswith(".pdf"):
        from pypdf import PdfReader  # type: ignore[import-untyped]

        reader = PdfReader(file_path)
        pages_text = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages_text).strip()

    if file_path.endswith(".docx"):
        import docx  # type: ignore[import-untyped]

        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    raise ValueError(f"Unsupported file type: {file_path}")


def ingest_files(file_paths: list[str]) -> list[PageContent]:
    results: list[PageContent] = []
    for path in file_paths:
        text = extract_text_from_file(path)
        content_hash = _compute_hash(text)
        # Derive a best-effort title from the filename
        title = path.split("/")[-1]
        results.append(
            PageContent(
                url=path,
                title=title,
                html="",
                text=text,
                content_hash=content_hash,
                source_type="file",
            )
        )
    return results
